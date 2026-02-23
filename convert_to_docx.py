#!/usr/bin/env python3
"""
Convert DARLINGTON_HND_SWE_REPORT.md to DOCX format using python-docx
"""

import os
import re
import sys
from pathlib import Path

# Check if python-docx is available, if not provide installation instructions
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("âŒ python-docx not installed.")
    print("Install with: pip install python-docx")
    sys.exit(1)

def parse_markdown_file(filepath):
    """Read and parse markdown file into structured content"""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    return content

def create_docx(markdown_content, output_path):
    """Convert markdown content to DOCX using python-docx"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title
    title = doc.add_heading('DARLINGTON HND SOFTWARE ENGINEERING REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.paragraph_format
    title_format.space_after = Pt(6)
    
    # Add metadata
    metadata = doc.add_paragraph()
    metadata.add_run('Author: ').bold = True
    metadata.add_run('Darlington\n')
    metadata.add_run('Date: ').bold = True
    metadata.add_run('December 15, 2025\n')
    metadata.add_run('Course: ').bold = True
    metadata.add_run('HND Software Engineering\n')
    metadata.add_run('Project: ').bold = True
    metadata.add_run('Accessibility Translator\n')
    metadata.add_run('Assessment Type: ').bold = True
    metadata.add_run('Major Project Report')
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Parse and add content sections
    lines = markdown_content.split('\n')
    in_code_block = False
    code_content = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip rendering note
        if 'Rendering Note:' in line:
            i += 1
            continue
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_content = []
            else:
                # End of code block
                in_code_block = False
                code_para = doc.add_paragraph('\n'.join(code_content))
                code_para_format = code_para.paragraph_format
                code_para_format.left_indent = Inches(0.5)
                code_para_format.right_indent = Inches(0.5)
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_content = []
        elif in_code_block:
            code_content.append(line)
        
        # Handle headings
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        
        # Handle horizontal rules
        elif line.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 80)
        
        # Handle image placeholders
        elif line.strip().startswith('!['):
            match = re.search(r'!\[([^\]]*)\]\(([^\)]*)\)', line)
            if match:
                alt_text = match.group(1)
                img_path = match.group(2)
                try:
                    if os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(5.5))
                        caption = doc.add_paragraph(alt_text)
                        caption_format = caption.paragraph_format
                        caption_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.style = 'Caption'
                except Exception as e:
                    doc.add_paragraph(f'[Image: {alt_text}]')
        
        # Handle lists
        elif line.strip().startswith('- '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.strip().startswith('* '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif re.match(r'^\d+\.\s', line.strip()):
            match = re.match(r'^(\d+)\.\s(.*)', line.strip())
            if match:
                doc.add_paragraph(match.group(2), style='List Number')
        
        # Handle tables
        elif line.strip().startswith('|'):
            # Parse table
            table_rows = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_rows.append(lines[i])
                i += 1
            
            if len(table_rows) > 2:  # Header + separator + at least one row
                # Parse header
                header = [cell.strip() for cell in table_rows[0].split('|')[1:-1]]
                
                # Create table
                table = doc.add_table(rows=1, cols=len(header))
                table.style = 'Light Grid Accent 1'
                
                # Add header
                header_cells = table.rows[0].cells
                for i, cell_text in enumerate(header):
                    header_cells[i].text = cell_text
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Add data rows
                for row_text in table_rows[2:]:
                    cells = [cell.strip() for cell in row_text.split('|')[1:-1]]
                    if cells:
                        row_cells = table.add_row().cells
                        for i, cell_text in enumerate(cells):
                            row_cells[i].text = cell_text
                i -= 1  # Adjust since we've incremented in the loop
        
        # Handle normal paragraphs
        elif line.strip() and not in_code_block:
            if not line.strip().startswith('---'):
                para = doc.add_paragraph(line.strip())
        
        i += 1
    
    # Save document
    doc.save(output_path)
    print(f"âœ… Document saved: {output_path}")

def main():
    """Main entry point"""
    markdown_file = 'DARLINGTON_HND_SWE_REPORT.md'
    output_file = 'DARLINGTON_HND_SWE_REPORT.docx'
    
    if not os.path.exists(markdown_file):
        print(f"âŒ File not found: {markdown_file}")
        sys.exit(1)
    
    print(f"ðŸ“„ Reading: {markdown_file}")
    content = parse_markdown_file(markdown_file)
    
    print(f"ðŸ”„ Converting to DOCX...")
    create_docx(content, output_file)
    
    if os.path.exists(output_file):
        size_kb = os.path.getsize(output_file) / 1024
        print(f"âœ… Conversion complete!")
        print(f"   File: {output_file}")
        print(f"   Size: {size_kb:.1f} KB")
    else:
        print(f"âŒ Conversion failed")
        sys.exit(1)

if __name__ == '__main__':
    main()

